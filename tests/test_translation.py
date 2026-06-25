import os
import sys
import docx
from docx import Document

# Add project root to sys.path so we can import docx_translator
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

def create_sample_docx(filename):
    doc = Document()
    doc.add_heading("Project Test Document", level=0)
    
    # Add a table
    table = doc.add_table(rows=3, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "ID"
    hdr_cells[1].text = "Source English"
    hdr_cells[2].text = "Translation" # Target column
    
    # Row 1
    row_cells = table.rows[1].cells
    row_cells[0].text = "1"
    row_cells[1].text = "This is a user manual for the localization software."
    row_cells[2].text = "This is a user manual for the localization software."
    
    # Row 2
    row_cells = table.rows[2].cells
    row_cells[0].text = "2"
    row_cells[1].text = "Save the output document to your desktop."
    row_cells[2].text = "Save the output document to your desktop."
    
    # Apply some basic formatting so we can verify layouts are retained
    for row in table.rows:
        row.height = docx.shared.Inches(0.4)
        for cell in row.cells:
            cell.width = docx.shared.Inches(2.0)
            
    doc.save(filename)
    print(f"Sample test document created at: {filename}")

def verify_output(filename):
    doc = Document(filename)
    table = doc.tables[0]
    
    print("\n--- Verifying Translated Content ---")
    for r_idx in range(1, len(table.rows)):
        row = table.rows[r_idx]
        print(f"Row {r_idx} Translated Text: {row.cells[2].text}")
        
        # Verify text was changed (e.g. from English to Spanish)
        if row.cells[2].text == row.cells[1].text:
            print("ERROR: Translation did not occur!")
            return False
            
    print("Verification Successful!")
    return True

def main():
    test_input = "test_sample.docx"
    test_output = "test_sample_es.docx"
    
    # Step 1: Create sample
    create_sample_docx(test_input)
    
    # Step 2: Initialize config and pipeline
    from docx_translator.config import Config
    from docx_translator.pipeline import LocalizationPipeline
    
    config = Config()
    config.settings["translation_backend"] = "google_translator"
    
    pipeline = LocalizationPipeline(
        config=config,
        target_lang="Spanish",
        source_lang="en",
        tm_path="test_tm.json",
        glossary_path="test_glossary.json"
    )
    
    # Mock translation calls to avoid network dependency in tests
    pipeline.translator.translate = lambda text, context=None: f"[ES] {text}"
    
    # Step 3: Run translation
    print("Running translation pipeline (ES)...")
    success = pipeline.translate_document(
        input_path=test_input,
        output_path=test_output,
        column_header="Translation"
    )
    
    if not success:
        print("Pipeline translation failed.")
        sys.exit(1)
        
    # Step 4: Verify
    if verify_output(test_output):
        print("All checks passed successfully.")
    else:
        print("Some checks failed.")
        sys.exit(1)
        
    # Cleanup temp test files
    for f in [test_input, test_output, "test_tm.json"]:
        if os.path.exists(f):
            os.remove(f)

if __name__ == "__main__":
    main()
