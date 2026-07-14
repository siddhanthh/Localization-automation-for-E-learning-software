import os
import re
import csv
import json
import time
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
from threading import Lock
import docx
from docx import Document

from docx_translator.translator import TranslatorBackend
from docx_translator.layout import restore_layouts

logger = logging.getLogger("docx_translator.pipeline")

class LocalizationPipeline:
    def __init__(self, config, target_lang="th", source_lang="en", tm_path=None, glossary_path=None):
        self.config = config
        self.target_lang = target_lang
        self.source_lang = source_lang
        self.translator = TranslatorBackend(config, target_lang=target_lang, source_lang=source_lang)
        
        self.tm_path = tm_path or "translation_memory.json"
        self.glossary_path = glossary_path or "glossary.json"
        
        self.tm = self._load_json_file(self.tm_path)
        self.glossary = self._load_json_file(self.glossary_path)
        self.tm_lock = Lock()

    def _load_json_file(self, path):
        if path and os.path.exists(path):
            try:
                with open(path, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception as e:
                logger.warning(f"Could not load JSON file {path}: {e}")
        return {}

    def _save_tm(self):
        if not self.tm_path:
            return
        with self.tm_lock:
            try:
                with open(self.tm_path, "w", encoding="utf-8") as f:
                    json.dump(self.tm, f, ensure_ascii=False, indent=2)
            except Exception as e:
                logger.error(f"Could not save translation memory to {self.tm_path}: {e}")

    def lookup_translation(self, text):
        """Lookup translation in glossary and translation memory (exact, normalized, case-insensitive)."""
        if not text:
            return None
        
        normalized_text = " ".join(text.split()).strip()
        
        # 1. Glossary Lookup (highest priority)
        if text in self.glossary:
            return self.glossary[text]
        if normalized_text in self.glossary:
            return self.glossary[normalized_text]
        
        # Case insensitive glossary check
        for k, v in self.glossary.items():
            if k.lower() == normalized_text.lower():
                return v

        # 2. TM Lookup
        if text in self.tm:
            return self.tm[text]["translation"]
        
        # Normalized or case differences in TM
        for k, entry in self.tm.items():
            k_norm = " ".join(k.split()).strip()
            if k_norm.lower() == normalized_text.lower():
                return entry["translation"]
                
        return None

    def add_to_tm(self, source, translation, context=None):
        with self.tm_lock:
            self.tm[source] = {
                "translation": translation,
                "context": context or "",
                "timestamp": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
                "status": "translated"
            }


    def scan_translation_cells(self, doc_or_path, column_header="Translation", limit_tables=None, limit_cells=None):
        """Finds all target cells in the specified column of all tables."""
        if isinstance(doc_or_path, str):
            doc = Document(doc_or_path)
        else:
            doc = doc_or_path
        cells_to_process = []

        
        table_count = 0
        cell_count = 0
        
        for t_idx, table in enumerate(doc.tables):
            if limit_tables and table_count >= limit_tables:
                break
                
            # Locate target column header
            target_col_idx = -1
            if len(table.rows) > 0:
                for col_idx, cell in enumerate(table.rows[0].cells):
                    if cell.text.strip() == column_header:
                        target_col_idx = col_idx
                        break
            
            if target_col_idx == -1:
                continue
                
            table_count += 1
            # Scan rows starting from row 1 to skip header
            for r_idx in range(1, len(table.rows)):
                if limit_cells and cell_count >= limit_cells:
                    break
                row = table.rows[r_idx]
                if target_col_idx < len(row.cells):
                    cell = row.cells[target_col_idx]
                    cells_to_process.append({
                        "table_index": t_idx,
                        "row_index": r_idx,
                        "cell_text": cell.text,
                        "header": column_header,
                        "cell": cell
                    })
                    cell_count += 1
                    
        return cells_to_process, table_count

    def analyze_document(self, input_path, column_header="Translation", report_path=None):
        """Scan document and generate analysis report."""
        logger.info(f"Scanning {input_path} for analysis...")
        cells, tables_found = self.scan_translation_cells(input_path, column_header)
        
        total_cells = len(cells)
        pre_translated_cells = 0
        protected_count = 0
        placeholders_count = 0
        translatable_cells = 0
        unique_strings = set()
        repeated_strings = {}
        
        for c in cells:
            text = c["cell_text"].strip()
            if not text:
                continue
            
            unique_strings.add(text)
            repeated_strings[text] = repeated_strings.get(text, 0) + 1
            
            if self.translator.should_skip_translation(text):
                # Check if it was because it already contains target characters or matches protected patterns
                # Let's count target script vs protected
                script_pattern = self.translator.should_skip_translation(text)
                protected_matched = False
                for pattern in re.compile(r'|'.join(re.escape(p) for p in re.compile(r'.*').findall(''))).findall(text):
                    pass # check manual patterns
                # Keep it simple for reporting
                protected_count += 1
                continue
                
            # Placeholders inside text
            placeholder_matches = re.findall(r'\{[^}]+\}|%s|%d|%[^ ]+s', text)
            if placeholder_matches:
                placeholders_count += len(placeholder_matches)
                
            translatable_cells += 1

        repeated_count = sum(1 for k, v in repeated_strings.items() if v > 1)

        report = {
            "input_file": input_path,
            "target_language": self.target_lang,
            "tables_containing_target_column": tables_found,
            "total_rows_scanned": len(cells),
            "cells_requiring_translation": translatable_cells,
            "skipped_or_protected_cells": protected_count,
            "placeholders_detected": placeholders_count,
            "repeated_strings_for_tm": repeated_count,
            "unique_translatable_strings": len(unique_strings) - protected_count
        }

        if report_path:
            with open(report_path, "w", encoding="utf-8") as f:
                json.dump(report, f, indent=2)
            logger.info(f"Analysis report saved to {report_path}")
            
        return report

    def _translate_with_runs(self, paragraph, context=None):
        """Processes paragraph runs with tag validation and fallback."""
        runs = paragraph.runs
        if not runs:
            return
            
        if len(runs) == 1:
            orig_text = runs[0].text
            cached = self.lookup_translation(orig_text)
            if cached:
                runs[0].text = cached
            else:
                try:
                    translated = self.translator.translate(orig_text, context)
                    self.add_to_tm(orig_text, translated, context)
                    runs[0].text = translated
                except Exception as e:
                    logger.error(f"Failed to translate run '{orig_text}': {e}")
            return

        # Tag-based translation for multiple runs
        tagged_parts = []
        for idx, run in enumerate(runs):
            tagged_parts.append(f"<r{idx}>{run.text}</r{idx}>")
            
        tagged_text = "".join(tagged_parts)
        
        try:
            translated_tagged = self.translator.translate(tagged_text, context)
            
            # Validate tags in response
            for idx in range(len(runs)):
                open_tag = f"<r{idx}>"
                close_tag = f"</r{idx}>"
                if open_tag not in translated_tagged or close_tag not in translated_tagged:
                    raise ValueError(f"Tag validation failed: Missing run tag r{idx}")
                    
            # Parse and apply translated text back to runs
            for idx, run in enumerate(runs):
                pattern = f"<r{idx}>(.*?)</r{idx}>"
                match = re.search(pattern, translated_tagged, re.DOTALL)
                if match:
                    run.text = match.group(1)
        except Exception as e:
            logger.warning(f"Tag-based translation failed ({e}). Falling back to run-by-run mode.")
            # Run-by-run fallback
            for run in runs:
                if run.text.strip():
                    cached = self.lookup_translation(run.text)
                    if cached:
                        run.text = cached
                    else:
                        try:
                            translated = self.translator.translate(run.text, context)
                            self.add_to_tm(run.text, translated, context)
                            run.text = translated
                        except Exception as ex:
                            logger.error(f"Failed to translate fallback run '{run.text}': {ex}")

    def translate_document(self, input_path, output_path, column_header="Translation", limit_tables=None, limit_cells=None, review_report_path=None, progress_callback=None):
        """Translate document cells, verify integrity, and restore layout in one unified step.
        
        Args:
            progress_callback: Optional callable(current, total, stage_text) for progress reporting.
        """
        doc = Document(input_path)
        cells, _ = self.scan_translation_cells(doc, column_header, limit_tables, limit_cells)
        
        total_cells = len(cells)
        # Total steps: cells for translation + cells for writing + 2 (verify + layout restore)
        total_steps = total_cells + total_cells + 2
        completed = [0]  # mutable counter for thread-safe incrementing
        
        def report_progress(stage_text):
            completed[0] += 1
            if progress_callback:
                progress_callback(completed[0], total_steps, stage_text)
        
        logger.info(f"Starting translation of {total_cells} cells...")
        
        # Concurrently fetch translations to pre-populate translation memory
        def process_cell(cell_info):
            text = cell_info["cell_text"]
            if self.translator.should_skip_translation(text):
                report_progress("Skipped (already in target language)")
                return
            
            # Check TM/glossary first
            if self.lookup_translation(text):
                report_progress("Found in translation memory")
                return
                
            context = f"Table header: {cell_info['header']}"
            try:
                translated = self.translator.translate(text, context)
                self.add_to_tm(text, translated, context)
                report_progress(f"Translated cell")
            except Exception as e:
                logger.error(f"Error translating cell: {e}")
                report_progress("Translation error")

        # Threaded run for performance
        num_workers = self.config.concurrency.get("num_workers", 8)
        if num_workers > 1 and len(cells) > 1:
            with ThreadPoolExecutor(max_workers=num_workers) as executor:
                futures = [executor.submit(process_cell, c) for c in cells]
                for future in as_completed(futures):
                    pass
        else:
            for c in cells:
                process_cell(c)
                
        # Save TM once at the end of translation block
        self._save_tm()
                
        # Copy original document to temp location for translating text runs
        temp_translated_path = output_path + ".tmp_trans"
        
        logger.info("Applying translations to document structure...")
        if progress_callback:
            progress_callback(completed[0], total_steps, "Applying translations to document...")
        
        # Scan again (using the already loaded `doc`) and apply run-by-run translation on target document using cached TM
        cells_to_write, _ = self.scan_translation_cells(doc, column_header, limit_tables, limit_cells)
        
        for c in cells_to_write:
            t_idx = c["table_index"]
            r_idx = c["row_index"]
            
            # Find the target cell in current document structure
            translation_col_idx = -1
            for idx, header_cell in enumerate(doc.tables[t_idx].rows[0].cells):
                if header_cell.text.strip() == column_header:
                    translation_col_idx = idx
                    break
                    
            if translation_col_idx != -1:
                target_cell = doc.tables[t_idx].rows[r_idx].cells[translation_col_idx]
                for paragraph in target_cell.paragraphs:
                    self._translate_with_runs(paragraph, context=f"Table: {t_idx}, Row: {r_idx}")
            report_progress(f"Applied translations to row {r_idx}")
                    
        doc.save(temp_translated_path)
        
        # Verify structure
        report_progress("Verifying document integrity...")
        if not self.verify_document_integrity(input_path, temp_translated_path):
            logger.error("Document structure check failed! Aborting translation.")
            if os.path.exists(temp_translated_path):
                os.remove(temp_translated_path)
            return False

        # Apply formatting and layouts to the translated document
        report_progress("Restoring document layouts...")
        logger.info("Finalizing document formatting...")
        restore_layouts(input_path, temp_translated_path, output_path)
        
        # Cleanup temp file
        if os.path.exists(temp_translated_path):
            os.remove(temp_translated_path)
            
        logger.info(f"Translation complete! Saved output to {output_path}")


        # Optional review report
        if review_report_path:
            self.generate_review_report(cells, review_report_path)

        return True

    def verify_document_integrity(self, original_path, target_path):
        """Checks structural constraints between original and translated files."""
        try:
            orig_doc = Document(original_path)
            target_doc = Document(target_path)
            
            if len(orig_doc.tables) != len(target_doc.tables):
                logger.error("Integrity check: Table count mismatch!")
                return False
                
            for idx in range(len(orig_doc.tables)):
                orig_table = orig_doc.tables[idx]
                target_table = target_doc.tables[idx]
                
                if len(orig_table.rows) != len(target_table.rows):
                    logger.error(f"Integrity check: Row count mismatch in Table {idx}!")
                    return False
                    
                if len(orig_table.columns) != len(target_table.columns):
                    logger.error(f"Integrity check: Column count mismatch in Table {idx}!")
                    return False
                    
            return True
        except Exception as e:
            logger.error(f"Integrity check failed to execute: {e}")
            return False

    def generate_review_report(self, cells, report_path):
        """Export localized review spreadsheet."""
        try:
            with open(report_path, "w", encoding="utf-8", newline="") as f:
                writer = csv.writer(f)
                writer.writerow(["Table Index", "Row Index", "Original", f"Translation ({self.target_lang})", "Status"])
                for c in cells:
                    orig = c["cell_text"]
                    trans = self.lookup_translation(orig) or orig
                    status = "Translated" if orig != trans else "Skipped/Protected"
                    writer.writerow([c["table_index"], c["row_index"], orig, trans, status])
            logger.info(f"Review report saved to {report_path}")
        except Exception as e:
            logger.warning(f"Could not generate review report: {e}")
